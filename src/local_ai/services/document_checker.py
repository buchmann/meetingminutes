"""Document checker: extract text from Word/PDF, improve via LLM (optionally in
the user's writing style), and regenerate as Word, PDF, or plain text.

Supported input formats:  .docx  .pdf  .txt  .md
Supported output formats: .docx  .pdf  .txt

The improvement step reuses :func:`local_ai.services.text_improver.improve_text`.
Longer documents are split at paragraph boundaries before improvement so that
each chunk fits the LLM context, then re-joined for output.
"""

from __future__ import annotations

import logging
from io import BytesIO
from pathlib import Path

from local_ai.services.text_improver import MAX_INPUT_CHARS, improve_text

logger = logging.getLogger(__name__)


# ── Text extraction ──────────────────────────────────────────────────────


def extract_text(file_path: Path, ext: str) -> str:
    """Pull the readable text out of a Word, PDF, or plain-text file.

    Returns a single string with paragraphs separated by blank lines so that
    the LLM (and downstream PDF/Word writers) can preserve structure.
    """
    ext = ext.lower().lstrip(".")
    if ext == "docx":
        return _extract_docx(file_path)
    if ext == "pdf":
        return _extract_pdf(file_path)
    if ext in {"txt", "md", "markdown"}:
        return file_path.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"Unsupported input format: .{ext}")


def _extract_docx(file_path: Path) -> str:
    from docx import Document  # python-docx

    doc = Document(str(file_path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
    # Tables: include cell text row-by-row so it isn't silently dropped.
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


# If the PDF text layer yields fewer than this many chars, treat it as a scan
# and fall back to OCR.
_OCR_MIN_CHARS = 200
_OCR_MAX_PAGES = 40


def _extract_pdf(file_path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    pages: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001
            logger.warning("PDF page extraction failed: %s", exc)
            continue
        # Normalize line breaks: PDF extractors often produce hard wraps mid-sentence.
        text = text.replace("\r\n", "\n").strip()
        if text:
            pages.append(text)
    layer_text = "\n\n".join(pages)

    # Scanned PDF (little/no text layer) → OCR fallback.
    if len(layer_text.strip()) < _OCR_MIN_CHARS:
        ocr_text = _ocr_pdf(file_path)
        if len(ocr_text.strip()) > len(layer_text.strip()):
            logger.info("PDF OCR fallback used (%d → %d chars)", len(layer_text), len(ocr_text))
            return ocr_text
    return layer_text


def _ocr_pdf(file_path: Path) -> str:
    """OCR a (scanned) PDF with Tesseract (German + English). Returns "" if the
    OCR toolchain (pdf2image/poppler/pytesseract/tesseract) is unavailable."""
    try:
        import pytesseract
        from pdf2image import convert_from_path
    except Exception as exc:  # noqa: BLE001
        logger.warning("OCR libs unavailable: %s", exc)
        return ""
    try:
        images = convert_from_path(str(file_path), dpi=200, last_page=_OCR_MAX_PAGES)
    except Exception as exc:  # noqa: BLE001
        logger.warning("pdf2image (poppler) failed: %s", exc)
        return ""
    out: list[str] = []
    for i, img in enumerate(images):
        try:
            txt = pytesseract.image_to_string(img, lang="deu+eng")
        except Exception as exc:  # noqa: BLE001
            logger.warning("Tesseract OCR failed on page %d: %s", i + 1, exc)
            continue
        txt = (txt or "").replace("\r\n", "\n").strip()
        if txt:
            out.append(txt)
    return "\n\n".join(out)


# ── Improvement (chunked when input exceeds the per-call limit) ──────────


def _split_into_chunks(text: str, max_chars: int) -> list[str]:
    """Split text at paragraph boundaries so no chunk exceeds *max_chars*.

    Paragraphs that are themselves longer than *max_chars* are hard-split on
    sentence-ish boundaries (". ", "! ", "? ") and finally on word boundaries.
    """
    if len(text) <= max_chars:
        return [text]

    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    buf = ""
    for para in paragraphs:
        if len(para) > max_chars:
            # Flush current buffer first
            if buf:
                chunks.append(buf)
                buf = ""
            # Hard-split the giant paragraph
            chunks.extend(_hard_split(para, max_chars))
            continue
        candidate = (buf + "\n\n" + para) if buf else para
        if len(candidate) > max_chars:
            chunks.append(buf)
            buf = para
        else:
            buf = candidate
    if buf:
        chunks.append(buf)
    return chunks


def _hard_split(text: str, max_chars: int) -> list[str]:
    out: list[str] = []
    remaining = text
    while len(remaining) > max_chars:
        # Prefer to break after a sentence terminator near the limit
        cut = remaining.rfind(". ", 0, max_chars)
        if cut < max_chars // 2:
            cut = remaining.rfind(" ", 0, max_chars)
        if cut <= 0:
            cut = max_chars
        out.append(remaining[: cut + 1].strip())
        remaining = remaining[cut + 1 :].lstrip()
    if remaining:
        out.append(remaining)
    return out


async def improve_document_text(
    text: str,
    *,
    style_profile: str | None,
    backend: str,
    openai_base_url: str,
    openai_api_key: str,
    openai_model: str,
    ollama_base_url: str,
    ollama_model: str,
    temperature: float | None = None,
) -> dict:
    """Run *text* through the text improver, splitting into chunks if needed.

    ``temperature`` is the "hallucination" dial (None → improver default 0.3).

    Returns a dict shaped like :func:`text_improver.improve_text` but with the
    full improved text and a summary of how many chunks were processed.
    """
    if not text.strip():
        return {
            "original": text,
            "improved": "",
            "language": "en",
            "chunks": 0,
            "error": "Document contained no readable text.",
        }

    # Use a slightly smaller-than-MAX chunk size to leave headroom for the
    # LLM prompt template that gets prepended to each chunk.
    chunk_budget = max(2000, MAX_INPUT_CHARS - 500)
    chunks = _split_into_chunks(text, chunk_budget)
    logger.info(
        "Document improvement: %d chars → %d chunk(s) of ~%d chars",
        len(text), len(chunks), chunk_budget,
    )

    improved_parts: list[str] = []
    detected_lang = "en"
    first_error: str | None = None

    for i, chunk in enumerate(chunks, start=1):
        logger.info("Improving chunk %d/%d (%d chars)", i, len(chunks), len(chunk))
        result = await improve_text(
            chunk,
            style_profile=style_profile,
            backend=backend,
            openai_base_url=openai_base_url,
            openai_api_key=openai_api_key,
            openai_model=openai_model,
            ollama_base_url=ollama_base_url,
            ollama_model=ollama_model,
            temperature=temperature,
        )
        if result.get("language"):
            detected_lang = result["language"]
        improved = (result.get("improved") or "").strip()
        if improved:
            improved_parts.append(improved)
        elif result.get("error"):
            # Keep the original chunk so output still has matching structure.
            improved_parts.append(chunk)
            if first_error is None:
                first_error = result["error"]

    return {
        "original": text,
        "improved": "\n\n".join(improved_parts),
        "language": detected_lang,
        "chunks": len(chunks),
        "error": first_error,
    }


# ── Output generation ────────────────────────────────────────────────────


def generate_markdown(text: str, *, title: str | None = None) -> bytes:
    """Render *text* as a Markdown document and return its bytes.

    The improved text is plain prose (paragraphs separated by blank lines), so
    the conversion is mostly pass-through — we just prepend a level-1 heading
    when a title is supplied and normalise paragraph spacing.
    """
    parts: list[str] = []
    if title:
        parts.append(f"# {title}\n")
    for block in text.split("\n\n"):
        block = block.strip()
        if block:
            parts.append(block)
    return ("\n\n".join(parts) + "\n").encode("utf-8")


def generate_docx(text: str, *, title: str | None = None) -> bytes:
    """Render *text* as a .docx and return its bytes.

    Paragraphs are split on blank lines; the optional *title* becomes a Heading-1.
    """
    from docx import Document  # python-docx
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if title:
        doc.add_heading(title, level=1)

    for block in text.split("\n\n"):
        block = block.strip()
        if block:
            doc.add_paragraph(block)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_pdf(text: str, *, title: str | None = None) -> bytes:
    """Render *text* as a PDF and return its bytes.

    Uses fpdf2 with a DejaVu Sans Unicode font when available (handles German
    umlauts and most Latin scripts). Falls back to the built-in Helvetica with
    a latin-1 best-effort encoding.
    """
    from fpdf import FPDF

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    # Try to register a Unicode font. fpdf2 ships with no fonts, but DejaVu Sans
    # is present in most Linux containers (and on macOS via Homebrew installs).
    font_loaded = False
    for path in (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
    ):
        if Path(path).exists():
            try:
                pdf.add_font("Doc", "", path)
                # Try bold variant of the same family if present.
                bold_path = path.replace("DejaVuSans.ttf", "DejaVuSans-Bold.ttf")
                if bold_path != path and Path(bold_path).exists():
                    pdf.add_font("Doc", "B", bold_path)
                font_loaded = True
                break
            except Exception as exc:  # noqa: BLE001
                logger.warning("Could not load font %s: %s", path, exc)

    base_font = "Doc" if font_loaded else "Helvetica"

    if title:
        try:
            pdf.set_font(base_font, "B", 16)
        except Exception:  # bold variant missing
            pdf.set_font(base_font, "", 16)
        pdf.multi_cell(0, 8, _safe_for_font(title, font_loaded))
        pdf.ln(2)

    pdf.set_font(base_font, "", 11)
    for block in text.split("\n\n"):
        block = block.strip()
        if block:
            pdf.multi_cell(0, 6, _safe_for_font(block, font_loaded))
            pdf.ln(2)

    out = pdf.output(dest="S")
    return bytes(out)


def _safe_for_font(text: str, unicode_font_loaded: bool) -> str:
    """When falling back to the built-in Helvetica (latin-1 only), drop chars
    that can't be encoded. With a Unicode font loaded, return the text unchanged.
    """
    if unicode_font_loaded:
        return text
    return text.encode("latin-1", "replace").decode("latin-1")
